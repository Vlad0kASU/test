import os
import docx
import pandas as pd
import xlsxwriter
from progress.bar import IncrementalBar


def main():
    directory = os.getcwd()
    files = os.listdir(directory)
    docs = list(filter(lambda x: x.endswith('.docx'), files))
    # print(images)
    table_for_block = [[], [], [], [], []]
    n=0
    bar = IncrementalBar('Прогресс', max=len(docs))
    for file in docs:
        doc = docx.Document(file)
        n+=1

        tables = []
        i = 0
        Text1 = []

        all_tables = doc.tables
        data_tables = {i: None for i in range(len(all_tables))}
        for i, table in enumerate(all_tables):
            # создаем список строк для таблицы `i` (пока пустые)
            data_tables[i] = [[] for _ in range(len(table.rows))]
            # проходимся по строкам таблицы `i`
            for j, row in enumerate(table.rows):
                # проходимся по ячейкам таблицы `i` и строки `j`
                for cell in row.cells:
                    # добавляем значение ячейки в соответствующий
                    # список, созданного словаря под данные таблиц
                    data_tables[i][j].append(cell.text)
                    Text1.append(cell.text)

        bar.next()
        print(f" Файл - {file}, Новое название - {Text1[5]}.docx")
        # Text2=[]
        # for para in doc.paragraphs:
        #     Text2.append(para.text)

        if not os.path.exists(f"{Text1[5]}.docx"):
            os.rename(file, f"{Text1[5]}.docx")
        # раздел 2
        # for j in range(len(data_tables[6])):
        #     table_for_block[0].append(Text1[5])
        #     for k in range(len(data_tables[6][j])):
        #         table_for_block[k+1].append(data_tables[6][j][k])

        # раздел 3

        # for i in range(len(Text1)):
        #     if Text1[i] == "Требования к образованию и обучению" and Text1[i - 22] == "Наименование" and Text1[i - 20] == "Код" and Text1[i - 18] == "Уровень квалификации":
        #         table_for_block[0].append(Text1[5])
        #         table_for_block[1].append(Text1[i-19])
        #         table_for_block[2].append(Text1[i-21])
        #         table_for_block[3].append(Text1[i-17])
        #         table_for_block[4].append(Text1[i+1].replace("\n", "\\"))

        # раздел Общие сведения
        # table_for_block[0].append(Text2[1])
        # table_for_block[1].append(Text1[1])
        # table_for_block[2].append(Text1[5])
        # table_for_block[3].append(Text1[3])
        # table_for_block[4].append(Text1[0].replace("\n"," "))



    # frame1 = {"1": Text1}
    # frame2 = {"2": Text2}
    # writer = pd.ExcelWriter("Text1.xlsx", engine="xlsxwriter")
    # dataframe1 = pd.DataFrame(frame1)
    # dataframe1.to_excel(writer,sheet_name="Лист1", index=False)
    # dataframe2 = pd.DataFrame(frame2)
    # dataframe2.to_excel(writer,sheet_name="Лист2", index=False)
    # writer.save()



    # name_block = ["Название","Регистрационный номер","Код","Наименование вида профессиональной деятельности","Утверждающий документ"]
    # dict_for_block = dict(zip(name_block, table_for_block))
    # data_frameBlock = pd.DataFrame(dict_for_block)
    # excelwriter = pd.ExcelWriter('ПрофстандартыОбщиеСведения.xlsx', engine='xlsxwriter')
    # data_frameBlock.to_excel(excelwriter, sheet_name='Лист1', index=False)
    # excelwriter.save()


if __name__ == '__main__':
    main()