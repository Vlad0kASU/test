import os
import docx
import pandas as pd
import xlsxwriter
from progress.bar import IncrementalBar


import sys


def excepthook(*args):
    import traceback
    import datetime
    import os
    try:
        exc_type, exc_val, exc_tb = args
    except:
        exc_type, exc_val, exc_tb = sys.exc_info()

    traceback.print_exception(exc_type, exc_val, exc_tb, file=sys.stdout)
    from asyncio import run
    async def send_me_massage(mes):
        from aiogram import Bot
        token = "5126991505:AAGbuAa-LtXn_AJinuCOC_x1IpwF4RlOsjc"
        bot = Bot(token)
        await bot.send_message(414337698, 'Раздел ЗУВ')
        await bot.send_message(414337698, mes)
        await bot.session.close()
    list_for_exc=traceback.format_exception(exc_type, exc_val, exc_tb)
    mes=""
    for i in range(len(list_for_exc)):
        mes+=list_for_exc[i]

    run(send_me_massage(mes))
    # with open(os.getcwd() + rf'\LOG{str(datetime.datetime.now())[:10]}.log', 'a+', encoding='UTF-8') as f:
    #     f.write(str(datetime.datetime.now()) + '\n')
    #     f.write(str(mes + '\n'))
    #     f.write('_' * 10 + '\n')
    # os.system('cls')
    input('программа завершена ошибкой')


def main():
    directory = os.getcwd()
    files = os.listdir(directory)
    docs = list(filter(lambda x: x.endswith('.docx'), files))

    n=0
    bar = IncrementalBar('Прогресс', max=len(docs))
    if not os.path.isdir(f"{directory}/ПС Раздел ЗУВ"):
        os.mkdir(f"{directory}/ПС Раздел ЗУВ")
    for file in docs:
        table_for_block = [[], [], [], []]
        doc = docx.Document(file)
        n+=1

        tables = []
        i = 0
        Text1 = []
        all_tables = doc.tables
        data_tables = {i: None for i in range(len(all_tables))}
        for i, table in enumerate(all_tables):
            # создаем список строк для таблицы `i` (пока пустые)
            data_tables[i] = [[] for _ in range(len(table.rows))]
            # проходимся по строкам таблицы `i`
            for j, row in enumerate(table.rows):
                # проходимся по ячейкам таблицы `i` и строки `j`
                for cell in row.cells:
                    # добавляем значение ячейки в соответствующий
                    # список, созданного словаря под данные таблиц
                    data_tables[i][j].append(cell.text)
                    Text1.append(cell.text)

        Text2 = []
        for para in doc.paragraphs:
            Text2.append(para.text)

        # раздел ЗУВ
        Numbers=[]
        for i in range(len(Text1)):
            if "/" in Text1[i] and Text1[i-1]=="Код":
                Numbers.append(i)

        regnumber = ""
        kodps = ""
        for i in range(100):
            if Text1[i] == "Регистрационный номер":
                regnumber = Text1[i - 2]
            if Text1[i] == "(наименование вида профессиональной деятельности)":
                kodps = Text1[i - 1]

        # print(len(Numbers))
        if len(Numbers) > 0:
            for i in range(1,len(Numbers)):
                for j in range(Numbers[i-1], Numbers[i]):
                    if Text1[j] == "Трудовые действия":
                        table_for_block[0].append(kodps)
                        table_for_block[1].append(Text1[Numbers[i-1]])
                        table_for_block[3].append(Text1[j+1])
                        table_for_block[2].append("1")
                    elif Text1[j] == "Необходимые умения":
                        table_for_block[0].append(kodps)
                        table_for_block[1].append(Text1[Numbers[i-1]])
                        table_for_block[3].append(Text1[j + 1])
                        table_for_block[2].append("2")
                    elif Text1[j] == "Необходимые знания":
                        table_for_block[0].append(kodps)
                        table_for_block[1].append(Text1[Numbers[i-1]])
                        table_for_block[3].append(Text1[j + 1])
                        table_for_block[2].append("3")
            for j in range(Numbers[len(Numbers)-1], len(Text1)):
                if Text1[j] == "Трудовые действия":
                    table_for_block[0].append(kodps)
                    table_for_block[1].append(Text1[Numbers[len(Numbers)-1]])
                    table_for_block[3].append(Text1[j+1])
                    table_for_block[2].append("1")
                elif Text1[j] == "Необходимые умения":
                    table_for_block[0].append(kodps)
                    table_for_block[1].append(Text1[Numbers[len(Numbers)-1]])
                    table_for_block[3].append(Text1[j + 1])
                    table_for_block[2].append("2")
                elif Text1[j] == "Необходимые знания":
                    table_for_block[0].append(kodps)
                    table_for_block[1].append(Text1[Numbers[len(Numbers)-1]])
                    table_for_block[3].append(Text1[j + 1])
                    table_for_block[2].append("3")
        else:
            table_for_block[0].append("Нет ТФ")
            table_for_block[1].append("")
            table_for_block[3].append("")
            table_for_block[2].append("")

        # for i in range(len(table_for_block)):
        #     print(len(table_for_block[i]))
        os.chdir(f"{directory}/ПС Раздел ЗУВ")
        name_block = ["Код ПС", "Код ТФ", "Категория ЗУВ",
                      "Наименование вида профессиональной деятельности"]
        dict_for_block = dict(zip(name_block, table_for_block))
        data_frameBlock = pd.DataFrame(dict_for_block)
        # txt = Text1[0].replace('\n', " ")
        # txt = txt[txt.find(" от "):]
        # txt = txt[1:len(txt) - 1]


        if not os.path.isfile(f"{kodps}_РегНомер_{regnumber}_{Text2[8]}..xlsx"):
            excelwriter = pd.ExcelWriter(f'{kodps}_РегНомер_{regnumber}_{Text2[8]}.xlsx', engine='xlsxwriter')
        else:
            excelwriter = pd.ExcelWriter(f'{kodps}_РегНомер_{regnumber}_{Text2[8]}_copy_{n}.xlsx', engine='xlsxwriter')

        data_frameBlock.to_excel(excelwriter, sheet_name='Лист1', index=False)
        excelwriter.save()

        os.chdir(f"{directory}")
        bar.next()
        print(f" Обработанный файл {file}")


if __name__ == '__main__':
    sys.excepthook = excepthook
    main()